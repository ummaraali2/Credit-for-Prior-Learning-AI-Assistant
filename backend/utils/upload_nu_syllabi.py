"""
Upload Northeastern Reference Syllabi to Milvus
These are REFERENCE documents (no student metadata)

Usage:
    python upload_nu_syllabi.py path/to/syllabus.pdf "INFO 5100"
    
Or upload multiple:
    python upload_nu_syllabi.py path/to/syllabi_folder/ "INFO 5100"

Or upload all PJM syllabi:
    python upload_nu_syllabi.py --all
"""

import os
import sys
import uuid
from pathlib import Path
from dotenv import load_dotenv
from ibm_watsonx_ai import APIClient, Credentials
from ibm_watsonx_ai.foundation_models.embeddings import Embeddings
from ibm_watsonx_ai.foundation_models.extensions.rag.vector_stores import MilvusVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import pdfplumber
import docx

load_dotenv()

# ==================== INITIALIZE SERVICES ====================

print("\n🔧 Initializing watsonx.ai services...\n")

# Initialize watsonx.ai
credentials = Credentials(
    api_key=os.getenv('WATSONX_AI_APIKEY'),
    url=os.getenv('WATSONX_AI_SERVICE_URL')
)
api_client = APIClient(credentials)
api_client.set.default_project(os.getenv('WATSONX_AI_PROJECT_ID'))

# Initialize Embeddings
embedding = Embeddings(
    model_id='ibm/slate-125m-english-rtrvr-v2',
    api_client=api_client
)

# Initialize Text Splitter
# Safe size: 1500 chars ≈ 375 tokens (well under 512 limit)
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1500,      # ← Safe size, under 512 token limit
    chunk_overlap=150,    # ← 10% overlap
    length_function=len,
    is_separator_regex=False,
)

# Initialize Milvus Vector Store
vector_store = MilvusVectorStore(
    api_client=api_client,
    connection_id=os.getenv('MILVUS_CONNECTION_ID'),
    collection_name='cpl_documents_v5',
    embedding_function=embedding
)

print("[SUCCESS] Services initialized")
print(f"   Collection: cpl_documents_v5 (L2 + HNSW)")
print(f"   Model: ibm/slate-125m-english-rtrvr-v2")
print(f"   Chunk size: 1500 chars (~375 tokens, safe limit)")
print(f"   Overlap: 150 chars\n")

# ==================== HELPER FUNCTIONS ====================

def extract_text(file_path):
    """
    Extract text from PDF/DOCX/TXT
    Uses pdfplumber for PDFs (handles font encoding issues better than PyPDF2)
    """
    try:
        if file_path.endswith('.pdf'):
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                raise ValueError("No text extracted from PDF. It may be scanned/image-based.")
            
            return text.strip()
        
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text.strip()
        
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        else:
            raise ValueError(f"Unsupported file type: {file_path}")
    except Exception as e:
        raise ValueError(f"Text extraction failed for {file_path}: {str(e)}")


def upload_nu_syllabus(file_path, course_code):
    """
    Upload a single NU reference syllabus
    
    Args:
        file_path: Path to syllabus file
        course_code: Course code (e.g., "INFO 5100")
    """
    
    filename = os.path.basename(file_path)
    document_id = str(uuid.uuid4())
    
    print(f"\n{'='*70}")
    print(f"[DOCUMENT] UPLOADING NU REFERENCE SYLLABUS")
    print(f"{'='*70}")
    print(f"   File: {filename}")
    print(f"   Course: {course_code}")
    print(f"   Document ID: {document_id}")
    print(f"   Type: nu_syllabus (NO student metadata)\n")
    
    # Extract text
    print("   [DOCUMENT] STEP 1: Extracting text...")
    text_content = extract_text(file_path)
    print(f"      [SUCCESS] Extracted {len(text_content)} characters")
    
    # Show preview of extracted text
    preview = text_content[:200].replace('\n', ' ')
    print(f"      📝 Preview: {preview}...\n")
    
    # Chunk document
    print(f"   [CHUNKING]  STEP 2: Chunking (size=1500, overlap=150)...")
    doc = Document(
        page_content=text_content,
        metadata={'document_name': filename}
    )
    chunks = text_splitter.split_documents([doc])
    print(f"      [SUCCESS] Created {len(chunks)} chunks\n")
    
    # Prepare documents for Milvus
    print("   [MILVUS] STEP 3: Preparing for Milvus...")
    documents = []
    char_position = 0
    
    for i, chunk in enumerate(chunks):
        chunk_text = chunk.page_content
        pk = f"{document_id}_{i}"
        
        documents.append({
            'content': chunk_text,  # ← CRITICAL: Must be 'content', not 'text'
            'metadata': {
                'pk': pk,
                'document_id': document_id,
                'document_name': filename,
                'document_type': 'nu_syllabus',
                'page': i + 1,
                'start_index': char_position,
                'sequence_number': i,
                'target_course': course_code,
                'student_name': '',
                'nuid': '',
                'request_type': ''
            }
        })
        
        char_position += len(chunk_text)
    
    print(f"      [SUCCESS] Prepared {len(documents)} documents\n")
    
    # Upload to Milvus
    print(f"   [UPLOADING] STEP 4: Uploading to Milvus...")
    result = vector_store.add_documents(documents)
    
    print(f"\n{'='*70}")
    print(f"[SUCCESS] UPLOAD COMPLETE!")
    print(f"{'='*70}")
    print(f"   Document: {filename}")
    print(f"   Course: {course_code}")
    print(f"   Chunks stored: {len(chunks)}")
    print(f"{'='*70}\n")
    
    return document_id


def upload_directory(directory_path, course_code):
    """Upload all PDFs/DOCX files in a directory"""
    
    path = Path(directory_path)
    
    if not path.exists():
        print(f"[ERROR] Directory not found: {directory_path}")
        return
    
    # Find all supported files
    files = list(path.glob('*.pdf')) + list(path.glob('*.docx')) + list(path.glob('*.txt'))
    
    if not files:
        print(f"[ERROR] No PDF/DOCX/TXT files found in: {directory_path}")
        return
    
    print(f"\n📂 Found {len(files)} file(s) in {directory_path}")
    print(f"   Course: {course_code}\n")
    
    success_count = 0
    fail_count = 0
    
    for file_path in files:
        try:
            upload_nu_syllabus(str(file_path), course_code)
            success_count += 1
        except Exception as e:
            print(f"[ERROR] Failed to upload {file_path.name}: {str(e)}\n")
            fail_count += 1
    
    print(f"\n{'='*70}")
    print(f"[ICEBERG] BATCH UPLOAD COMPLETE")
    print(f"{'='*70}")
    print(f"   [SUCCESS] Success: {success_count}")
    print(f"   [ERROR] Failed: {fail_count}")
    print(f"   📂 Total: {len(files)}")
    print(f"{'='*70}\n")


def upload_all_pjm_syllabi():
    """Upload all PJM syllabi from PJMSyllabi-Cleaned folder"""
    
    syllabi_folder = "PJMSyllabi-Cleaned"
    
    if not os.path.exists(syllabi_folder):
        print(f"[ERROR] Folder not found: {syllabi_folder}")
        return
    
    # Find all PJM*.txt files
    files = sorted([f for f in os.listdir(syllabi_folder) if f.startswith("PJM") and f.endswith(".txt")])
    
    if not files:
        print(f"[ERROR] No PJM*.txt files found in {syllabi_folder}")
        return
    
    print(f"\n{'='*70}")
    print(f"🎓 UPLOADING ALL PJM SYLLABI")
    print(f"{'='*70}")
    print(f"   Found {len(files)} files:")
    for f in files:
        print(f"      - {f}")
    print(f"   Chunk size: 1500 chars (safe, under token limit)")
    print(f"   Collection: cpl_documents_v5")
    print(f"{'='*70}\n")
    
    success_count = 0
    fail_count = 0
    
    for filename in files:
        file_path = os.path.join(syllabi_folder, filename)
        # Extract course code from filename (e.g., PJM5900.txt -> PJM5900)
        course_code = filename.replace(".txt", "").replace("_", "")
        
        try:
            upload_nu_syllabus(file_path, course_code)
            success_count += 1
        except Exception as e:
            print(f"[ERROR] Failed to upload {filename}: {str(e)}\n")
            fail_count += 1
    
    print(f"\n{'='*70}")
    print(f"🎉 ALL PJM SYLLABI UPLOADED!")
    print(f"{'='*70}")
    print(f"   [SUCCESS] Success: {success_count}")
    print(f"   [ERROR] Failed: {fail_count}")
    print(f"   📂 Total: {len(files)}")
    print(f"{'='*70}\n")


# ==================== MAIN ====================

if __name__ == '__main__':
    
    # Check for --all flag
    if len(sys.argv) == 2 and sys.argv[1] == '--all':
        upload_all_pjm_syllabi()
        sys.exit(0)
    
    if len(sys.argv) < 3:
        print("\n[ERROR] Usage Error\n")
        print("Upload single file:")
        print("   python upload_nu_syllabi.py path/to/syllabus.pdf 'PJM5900'\n")
        print("Upload directory:")
        print("   python upload_nu_syllabi.py path/to/syllabi_folder/ 'PJM5900'\n")
        print("Upload ALL PJM syllabi from PJMSyllabi-Cleaned folder:")
        print("   python upload_nu_syllabi.py --all\n")
        sys.exit(1)
    
    path_arg = sys.argv[1]
    course_code = sys.argv[2]
    
    print(f"\n{'='*70}")
    print(f"🎓 NU REFERENCE SYLLABUS UPLOADER")
    print(f"{'='*70}")
    print(f"   Target: {path_arg}")
    print(f"   Course: {course_code}")
    print(f"   Collection: cpl_documents_v5 (L2 + HNSW)")
    print(f"   Chunk size: 1500 chars")
    print(f"{'='*70}\n")
    
    # Check if path is file or directory
    path = Path(path_arg)
    
    if path.is_file():
        try:
            upload_nu_syllabus(str(path), course_code)
            print("[SUCCESS] All done!\n")
        except Exception as e:
            print(f"[ERROR] Upload failed: {str(e)}\n")
            sys.exit(1)
    
    elif path.is_dir():
        upload_directory(str(path), course_code)
        print("[SUCCESS] All done!\n")
    
    else:
        print(f"[ERROR] Path not found: {path_arg}\n")
        sys.exit(1)