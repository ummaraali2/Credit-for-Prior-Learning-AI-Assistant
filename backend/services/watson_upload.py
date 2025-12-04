# CPL Upload Service

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
from dotenv import load_dotenv
from ibm_watsonx_ai import APIClient, Credentials
from ibm_watsonx_ai.foundation_models.embeddings import Embeddings
from ibm_watsonx_ai.foundation_models.extensions.rag.vector_stores import MilvusVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import PyPDF2
import docx
import io
import uuid
from datetime import datetime
import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))
from handlers.iceberg_handler import get_iceberg_handler
from handlers.cos_handler import get_cos_handler

load_dotenv()
app = Flask(__name__)
CORS(app)

# Configuration
CHUNK_SIZE = 800
CHUNK_OVERLAP = 150

# Initialize services
credentials = Credentials(
    api_key=os.getenv('WATSONX_AI_APIKEY'),
    url=os.getenv('WATSONX_AI_SERVICE_URL')
)
api_client = APIClient(credentials)
api_client.set.default_project(os.getenv('WATSONX_AI_PROJECT_ID'))
embedding = Embeddings(
    model_id='ibm/slate-125m-english-rtrvr-v2',
    api_client=api_client
)

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    length_function=len,
    is_separator_regex=False,
)

vector_store = MilvusVectorStore(
    api_client=api_client,
    connection_id=os.getenv('MILVUS_CONNECTION_ID'),
    collection_name='cpl_documents_v5',
    embedding_function=embedding
)

iceberg = get_iceberg_handler()

cos = get_cos_handler()

# Helper functions

def extract_text(file_bytes, filename):
    try:
        if filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()

        elif filename.endswith('.docx'):
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([para.text for para in doc.paragraphs])
            return text.strip()

        elif filename.endswith('.txt'):
            return file_bytes.decode('utf-8')

        else:
            raise ValueError("Unsupported file type")
    except Exception as e:
        raise ValueError(f"Text extraction failed: {str(e)}")

def create_embedded_content(chunk_text, student_name, nuid, document_type, request_type, target_course, filename):
    enriched_content = f"""[DOCUMENT METADATA]
NUID: {nuid}
Student Name: {student_name}
Document Type: {document_type}
Request Type: {request_type}
Target Course: {target_course}
Source File: {filename}
[END METADATA]

[CONTENT]
{chunk_text}
[END CONTENT]"""

    return enriched_content

def safe_truncate_content(enriched_content, max_tokens=450):
    max_chars = int(max_tokens / 0.75)

    if len(enriched_content) > max_chars:
        print(f"      [WARNING]  Truncating chunk from {len(enriched_content)} to {max_chars} chars")
        truncated = enriched_content[:max_chars]
        last_period = truncated.rfind('.')
        if last_period > max_chars * 0.8:
            truncated = truncated[:last_period + 1]
        return truncated

    return enriched_content

@app.route('/api/upload-to-watsonx', methods=['POST'])
def upload_to_watsonx():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['file']
        filename = file.filename
        file_bytes = file.read()
        student_name = request.form.get('studentName', 'Unknown')
        nuid = request.form.get('nuid', 'N/A')
        request_type = request.form.get('requestType', 'Not Specified')
        target_course = request.form.get('targetCourse', 'Not Specified')
        document_id = str(uuid.uuid4())

        # Determine document type
        filename_lower = filename.lower()
        if 'transcript' in filename_lower:
            document_type = 'transcript'
        elif 'resume' in filename_lower or 'cv' in filename_lower:
            document_type = 'resume'
        else:
            document_type = 'student_syllabus'

        

        
        # Extract text
        text_content = extract_text(file_bytes, filename)

        # Chunk document
        doc = Document(
            page_content=text_content,
            metadata={'document_name': filename}
        )
        chunks = text_splitter.split_documents([doc])

        # Prepare documents for Milvus
        documents = []
        char_position = 0
        truncated_count = 0

        for i, chunk in enumerate(chunks):
            chunk_text = chunk.page_content

            # Create content with EMBEDDED metadata
            enriched_content = create_embedded_content(
                chunk_text=chunk_text,
                student_name=student_name,
                nuid=nuid,
                document_type=document_type,
                request_type=request_type,
                target_course=target_course,
                filename=filename
            )

            # Safety truncation to stay under 512 token limit
            original_len = len(enriched_content)
            enriched_content = safe_truncate_content(enriched_content, max_tokens=450)
            if len(enriched_content) < original_len:
                truncated_count += 1

            # Use 'content' key (watsonx.ai SDK maps this to 'text' field)
            documents.append({
                'content': enriched_content,  # Embedded metadata + safety truncated
                'metadata': {
                    'document_id': document_id,
                    'document_name': filename,
                    'document_type': document_type,
                    'page': i + 1,
                    'start_index': char_position,
                    'sequence_number': i,
                    # Student context (also in separate fields for filtering)
                    'student_name': student_name,
                    'nuid': nuid,
                    'target_course': target_course,
                    'request_type': request_type
                }
            })

            char_position += len(chunk_text)

        if truncated_count > 0:

        # Show sample
        if documents:
            print(f"\n   [REQUEST] Sample chunk preview:")
            sample = documents[0]['content'][:350].replace('\n', '\n   ')

        # Upload to Milvus
        print(f"\n   [UPLOADING] STEP 4: Uploading to Milvus...")
        result = vector_store.add_documents(documents)

        
        print("\n   [COS]  PART 2: Storing in COS...")

        try:
            cos_key = cos.upload_document(
                file_bytes=file_bytes,
                document_id=document_id,
                filename=filename,
                metadata={
                    'student_name': student_name,
                    'nuid': nuid,
                    'request_type': request_type,
                    'target_course': target_course
                }
            )
        except Exception as cos_error:
            cos_key = None

        
        print("\n   [ICEBERG] PART 3: Storing in ICEBERG...")

        request_id = iceberg.insert_request({
            'document_id': document_id,
            'student_name': student_name,
            'nuid': nuid,
            'request_type': request_type,
            'target_course': target_course,
            'document_name': filename,
            'cos_key': cos_key  # Store COS reference
        })

        if request_id:
        else:

        # ==================== COMPLETE ====================

        print(f"\n   [SUCCESS] UPLOAD COMPLETE!")

        return jsonify({
            'success': True,
            'document_id': document_id,
            'request_id': request_id,
            'filename': filename,
            'document_type': document_type,
            'student_name': student_name,
            'nuid': nuid,
            'request_type': request_type,
            'target_course': target_course,
            'chunks_created': len(chunks),
            'chunks_truncated': truncated_count,
            'chunk_size': CHUNK_SIZE,
            'characters_processed': len(text_content),
            'metadata_embedded': True,
            'cos_key': cos_key,
            'storage': {
                'milvus': f'{len(chunks)} chunks (embedded metadata, token-safe)',
                'cos': f'Original file stored: {cos_key}' if cos_key else 'COS upload failed',
                'iceberg': 'Student metadata stored'
            }
        })

    except Exception as e:
        print(f"\n[ERROR] ========== ERROR ==========")
        print(f"File: {filename if 'filename' in locals() else 'Unknown'}")
        print(f"Error: {str(e)}")
        print(f"=============================\n")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500
@app.route('/api/download-document/<document_id>/<filename>', methods=['GET'])
def download_document(document_id, filename):
    """
    Download original document from COS
    """
    try:
        print(f"\n[RECEIVED] Downloading document: {document_id}/{filename}")

        # Retrieve from COS
        file_bytes, metadata = cos.get_document_by_id(document_id, filename)

        # Send file to user
        return send_file(
            io.BytesIO(file_bytes),
            as_attachment=True,
            download_name=filename,
            mimetype='application/octet-stream'
        )

    except Exception as e:
        print(f"[ERROR] Download error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
@app.route('/api/preview-document/<document_id>/<filename>', methods=['GET'])
def preview_document(document_id, filename):
    """
    Preview document inline in browser (instead of downloading)
    """
    try:
        print(f"\n[PREVIEW] Previewing document: {document_id}/{filename}")

        # Retrieve from COS
        file_bytes, metadata = cos.get_document_by_id(document_id, filename)

        # Determine MIME type
        if filename.lower().endswith('.pdf'):
            mimetype = 'application/pdf'
        elif filename.lower().endswith(('.jpg', '.jpeg')):
            mimetype = 'image/jpeg'
        elif filename.lower().endswith('.png'):
            mimetype = 'image/png'
        elif filename.lower().endswith('.docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif filename.lower().endswith('.doc'):
            mimetype = 'application/msword'
        else:
            mimetype = 'application/octet-stream'

        # Send file for inline display
        return send_file(
            io.BytesIO(file_bytes),
            as_attachment=False,  # Display inline, not download
            download_name=filename,
            mimetype=mimetype
        )

    except Exception as e:
        print(f"[ERROR] Preview error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
@app.route('/api/view-document/<document_id>/<filename>', methods=['GET'])
def view_document(document_id, filename):
    """
    View document metadata from COS (without downloading)
    """
    try:
        file_bytes, metadata = cos.get_document_by_id(document_id, filename)

        return jsonify({
            'success': True,
            'document_id': document_id,
            'filename': filename,
            'size': len(file_bytes),
            'metadata': metadata
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
@app.route('/api/get-requests', methods=['GET'])
def get_requests():
    """Get all CPL requests FROM ICEBERG TABLE"""
    try:
        print("\n[QUERY] ========== QUERYING ICEBERG FOR REQUESTS ==========")
        requests = iceberg.get_all_requests()

        return jsonify({
            'success': True,
            'requests': requests,
            'count': len(requests),
            'source': 'iceberg'
        })

    except Exception as e:
        print(f"\n[ERROR] ICEBERG QUERY ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500
@app.route('/api/update-status', methods=['PUT'])
def update_status():
    """Update request status in Iceberg table"""
    try:
        data = request.json
        success = iceberg.update_status(
            request_id=data.get('requestId'),
            status=data.get('status'),
            credits=data.get('credits'),
            notes=data.get('notes', ''),
            updated_by=data.get('updatedBy', 'Advisor')
        )

        if success:
            return jsonify({'success': True})
        else:
            return jsonify({'success': False}), 500

    except Exception as e:
        print(f"[ERROR] Status update error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
@app.route('/api/search', methods=['POST'])
def search_documents():
    """Search documents in Milvus vector store"""
    try:
        data = request.json
        query = data.get('query')
        top_k = data.get('topK', 5)

        if not query:
            return jsonify({'success': False, 'error': 'Query required'}), 400

        results = vector_store.search(query, k=top_k)

        return jsonify({
            'success': True,
            'query': query,
            'results': results,
            'count': len(results)
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
@app.route('/health', methods=['GET'])
def health():
    """Health check"""
    return jsonify({
        'status': 'OK',
        'service': 'watsonx.ai Upload Service',
        'configuration': {
            'embedding_model': 'ibm/slate-125m-english-rtrvr-v2',
            'token_limit': 512,
            'chunk_size': CHUNK_SIZE,
            'chunk_overlap': CHUNK_OVERLAP,
            'milvus_collection': 'cpl_documents_v5',
            'cos_bucket': os.getenv('COS_BUCKET_NAME', 'cpl-documents'),
            'metadata_embedded': True,
            'safety_truncation': True,
            'cos_enabled': True
        }
    })
@app.route('/', methods=['GET'])
def home():
    """API information"""
    return jsonify({
        'service': 'watsonx.ai Upload Service',
        'version': '7.0 - COS Integration + Token Fix',
        'features': [
            f'Optimized chunks ({CHUNK_SIZE} chars)',
            'Metadata embedded in content',
            'Safety truncation for 512 token limit',
            'COS storage for original files',
            'Download original documents',
            'Works with Prompt Lab vector search'
        ],
        'endpoints': {
            'upload': 'POST /api/upload-to-watsonx',
            'download': 'GET /api/download-document/<document_id>/<filename>',
            'view': 'GET /api/view-document/<document_id>/<filename>',
            'get_requests': 'GET /api/get-requests',
            'update_status': 'PUT /api/update-status',
            'search': 'POST /api/search'
        }
    })
# ==================== START SERVER ====================

if __name__ == '__main__':
    print("\n[UPLOADING] ========== STARTING SERVER ==========")
    print(f"Service: watsonx.ai Upload Service v7.0")
    print(f"Port: 5000")
    print(f"Collection: cpl_documents_v5")
    print(f"COS Bucket: {os.getenv('COS_BUCKET_NAME', 'cpl-documents')}")
    print(f"Chunk size: {CHUNK_SIZE} (with {CHUNK_OVERLAP} overlap)")
    print(f"Token limit: 512 (with safety truncation at 450)")
    print(f"Features: Embedded metadata + COS storage + token safety")
    print("======================================\n")
    app.run(host='0.0.0.0', port=5000, debug=True)